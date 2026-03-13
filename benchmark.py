"""
MongoDB Benchmark: Azure Cosmos DB (DocumentDB) vs MongoDB Atlas
Compares Insert, Update, Aggregation, Lookup, and Delete operations.
"""

import os
import time
import random
import string
import statistics
from pymongo import MongoClient
from pymongo.errors import ConnectionFailure, OperationFailure
from datetime import datetime
from tabulate import tabulate
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv("config.env")

# ---------------------------------------------------------------------------
# Connection strings (loaded from config.env)
# ---------------------------------------------------------------------------
COSMOSDB_URI = os.getenv("COSMOSDB_URI")
ATLAS_URI = os.getenv("ATLAS_URI")

DB_NAME = "benchmark_db"
USERS_COLLECTION = "users"
ORDERS_COLLECTION = "orders"

NUM_USERS = 200
ORDERS_PER_USER = 50
ITERATIONS = 5  # repeat each operation for stability


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def random_string(length: int) -> str:
    return "".join(random.choices(string.ascii_letters + string.digits, k=length))


def generate_user_doc(user_id: int) -> dict:
    """Generate a ~2 KB user document."""
    return {
        "user_id": user_id,
        "first_name": random_string(15),
        "last_name": random_string(15),
        "email": f"user{user_id}@example.com",
        "age": random.randint(18, 70),
        "address": {
            "street": random_string(30),
            "city": random_string(15),
            "state": random_string(10),
            "zip": random_string(5),
            "country": random_string(10),
        },
        "phone": random_string(12),
        "bio": random_string(800),          # bulk of the 2 KB
        "preferences": {
            "language": random.choice(["en", "fr", "es", "de"]),
            "theme": random.choice(["dark", "light"]),
            "notifications": random.choice([True, False]),
        },
        "tags": [random_string(8) for _ in range(10)],
        "padding": random_string(500),       # pad to ~2 KB
    }


def generate_order_docs(user_id: int, count: int) -> list[dict]:
    """Generate *count* order documents for a given user."""
    orders = []
    for i in range(count):
        order = {
            "order_id": user_id * ORDERS_PER_USER + i,
            "user_id": user_id,
            "status": random.choice(["pending", "shipped", "delivered", "cancelled"]),
            "total_amount": round(random.uniform(10, 500), 2),
            "currency": "USD",
            "items": [
                {
                    "product_name": random_string(20),
                    "quantity": random.randint(1, 5),
                    "unit_price": round(random.uniform(5, 100), 2),
                }
                for _ in range(random.randint(1, 4))
            ],
            "shipping_address": {
                "street": random_string(30),
                "city": random_string(15),
                "state": random_string(10),
                "zip": random_string(5),
            },
            "created_at": f"2025-{random.randint(1,12):02d}-{random.randint(1,28):02d}",
        }
        orders.append(order)
    return orders


# ---------------------------------------------------------------------------
# Benchmark helpers
# ---------------------------------------------------------------------------
def time_operation(func, *args, **kwargs):
    """Return elapsed time in ms and the function result."""
    start = time.perf_counter()
    result = func(*args, **kwargs)
    elapsed_ms = (time.perf_counter() - start) * 1000
    return elapsed_ms, result


def cleanup(db):
    """Drop benchmark collections."""
    db.drop_collection(USERS_COLLECTION)
    db.drop_collection(ORDERS_COLLECTION)


# ---------------------------------------------------------------------------
# Individual benchmark functions
# ---------------------------------------------------------------------------
def bench_insert(db):
    """Bulk-insert users and orders. Returns total ms."""
    users = [generate_user_doc(i) for i in range(NUM_USERS)]
    orders = []
    for i in range(NUM_USERS):
        orders.extend(generate_order_docs(i, ORDERS_PER_USER))

    ms_users, _ = time_operation(db[USERS_COLLECTION].insert_many, users)

    # Insert orders in batches to avoid connection timeouts on large payloads
    BATCH_SIZE = 200
    ms_orders = 0
    for start in range(0, len(orders), BATCH_SIZE):
        batch = orders[start : start + BATCH_SIZE]
        ms_batch, _ = time_operation(db[ORDERS_COLLECTION].insert_many, batch)
        ms_orders += ms_batch

    return ms_users + ms_orders


def bench_update_inc(db):
    """Update all users: increment age by 1."""
    ms, _ = time_operation(
        db[USERS_COLLECTION].update_many,
        {},
        {"$inc": {"age": 1}},
    )
    return ms


def bench_update_set_nested(db):
    """Update all users: set nested address.city and a new flag."""
    ms, _ = time_operation(
        db[USERS_COLLECTION].update_many,
        {},
        {"$set": {"address.city": "BenchmarkCity", "updated": True}},
    )
    return ms


def bench_update_push_array(db):
    """Update all users: push a new tag into the tags array."""
    ms, _ = time_operation(
        db[USERS_COLLECTION].update_many,
        {},
        {"$push": {"tags": "bench_tag"}},
    )
    return ms


def bench_update_conditional(db):
    """Conditional update: set status='vip' for users older than 40."""
    ms, _ = time_operation(
        db[USERS_COLLECTION].update_many,
        {"age": {"$gt": 40}},
        {"$set": {"status": "vip"}},
    )
    return ms


def bench_update_orders_status(db):
    """Update orders: change all 'pending' orders to 'processing'."""
    ms, _ = time_operation(
        db[ORDERS_COLLECTION].update_many,
        {"status": "pending"},
        {"$set": {"status": "processing"}},
    )
    return ms


# -- Find operations --------------------------------------------------------

def bench_find_by_id(db):
    """Find a single user by user_id (indexed)."""
    target = random.randint(0, NUM_USERS - 1)
    ms, _ = time_operation(
        lambda: db[USERS_COLLECTION].find_one({"user_id": target})
    )
    return ms


def bench_find_range(db):
    """Find users with age in a range."""
    ms, _ = time_operation(
        lambda: list(db[USERS_COLLECTION].find({"age": {"$gte": 30, "$lte": 50}}))
    )
    return ms


def bench_find_regex(db):
    """Find users whose email starts with 'user1'."""
    ms, _ = time_operation(
        lambda: list(db[USERS_COLLECTION].find({"email": {"$regex": "^user1"}}))
    )
    return ms


def bench_find_projection(db):
    """Find all users but only return user_id and first_name (projection)."""
    ms, _ = time_operation(
        lambda: list(db[USERS_COLLECTION].find({}, {"user_id": 1, "first_name": 1, "_id": 0}))
    )
    return ms


def bench_find_sort_limit(db):
    """Find top 20 users sorted by age descending."""
    ms, _ = time_operation(
        lambda: list(db[USERS_COLLECTION].find().sort("age", -1).limit(20))
    )
    return ms


def bench_find_orders_for_user(db):
    """Find all orders for a specific user (indexed field)."""
    target = random.randint(0, NUM_USERS - 1)
    ms, _ = time_operation(
        lambda: list(db[ORDERS_COLLECTION].find({"user_id": target}))
    )
    return ms


def bench_find_multi_field(db):
    """Find users by age range AND preference language (compound filter)."""
    ms, _ = time_operation(
        lambda: list(db[USERS_COLLECTION].find({
            "age": {"$gte": 25, "$lte": 45},
            "preferences.language": "fr",
        }))
    )
    return ms


def bench_find_in_array(db):
    """Find users whose tags array contains a specific value."""
    ms, _ = time_operation(
        lambda: list(db[USERS_COLLECTION].find({"tags": "bench_tag"}))
    )
    return ms


def bench_find_orders_by_amount(db):
    """Find orders with total_amount > 300 sorted by amount desc, limit 50."""
    ms, _ = time_operation(
        lambda: list(
            db[ORDERS_COLLECTION]
            .find({"total_amount": {"$gt": 300}})
            .sort("total_amount", -1)
            .limit(50)
        )
    )
    return ms


def bench_find_orders_status_user(db):
    """Find delivered orders for a specific user (compound: user_id + status)."""
    target = random.randint(0, NUM_USERS - 1)
    ms, _ = time_operation(
        lambda: list(db[ORDERS_COLLECTION].find({
            "user_id": target,
            "status": "delivered",
        }))
    )
    return ms


def bench_find_count(db):
    """Count orders with status 'shipped'."""
    ms, _ = time_operation(
        lambda: db[ORDERS_COLLECTION].count_documents({"status": "shipped"})
    )
    return ms


def bench_find_distinct(db):
    """Get distinct order statuses."""
    ms, _ = time_operation(
        lambda: db[ORDERS_COLLECTION].distinct("status")
    )
    return ms


# -- Aggregation operations -------------------------------------------------

def bench_agg_group_status(db):
    """Aggregate orders: group by status, compute count + total + avg."""
    pipeline = [
        {"$group": {
            "_id": "$status",
            "count": {"$sum": 1},
            "total_amount": {"$sum": "$total_amount"},
            "avg_amount": {"$avg": "$total_amount"},
        }},
        {"$sort": {"count": -1}},
    ]
    ms, _ = time_operation(lambda: list(db[ORDERS_COLLECTION].aggregate(pipeline)))
    return ms


def bench_agg_unwind_items(db):
    """Unwind order items and compute the most popular product names."""
    pipeline = [
        {"$unwind": "$items"},
        {"$group": {
            "_id": "$items.product_name",
            "times_ordered": {"$sum": "$items.quantity"},
        }},
        {"$sort": {"times_ordered": -1}},
        {"$limit": 10},
    ]
    ms, _ = time_operation(lambda: list(db[ORDERS_COLLECTION].aggregate(pipeline)))
    return ms


def bench_agg_bucket_amount(db):
    """Bucket orders by total_amount ranges."""
    pipeline = [
        {"$bucket": {
            "groupBy": "$total_amount",
            "boundaries": [0, 50, 100, 200, 300, 500, 1000],
            "default": "Other",
            "output": {
                "count": {"$sum": 1},
                "avg_amount": {"$avg": "$total_amount"},
            },
        }},
    ]
    ms, _ = time_operation(lambda: list(db[ORDERS_COLLECTION].aggregate(pipeline)))
    return ms


def bench_agg_user_order_stats(db):
    """Per-user aggregation: count orders and total spend (on orders collection)."""
    pipeline = [
        {"$group": {
            "_id": "$user_id",
            "order_count": {"$sum": 1},
            "total_spent": {"$sum": "$total_amount"},
            "max_order": {"$max": "$total_amount"},
            "min_order": {"$min": "$total_amount"},
        }},
        {"$sort": {"total_spent": -1}},
        {"$limit": 20},
    ]
    ms, _ = time_operation(lambda: list(db[ORDERS_COLLECTION].aggregate(pipeline)))
    return ms


def bench_agg_date_breakdown(db):
    """Group orders by month from created_at string."""
    pipeline = [
        {"$addFields": {
            "month": {"$substr": ["$created_at", 5, 2]},
        }},
        {"$group": {
            "_id": "$month",
            "count": {"$sum": 1},
            "revenue": {"$sum": "$total_amount"},
        }},
        {"$sort": {"_id": 1}},
    ]
    ms, _ = time_operation(lambda: list(db[ORDERS_COLLECTION].aggregate(pipeline)))
    return ms


def bench_lookup(db):
    """$lookup: join users → orders, project summary. Returns ms."""
    pipeline = [
        {"$lookup": {
            "from": ORDERS_COLLECTION,
            "localField": "user_id",
            "foreignField": "user_id",
            "as": "user_orders",
        }},
        {"$project": {
            "user_id": 1,
            "first_name": 1,
            "last_name": 1,
            "order_count": {"$size": "$user_orders"},
            "total_spent": {"$sum": "$user_orders.total_amount"},
        }},
        {"$sort": {"total_spent": -1}},
        {"$limit": 10},
    ]
    ms, _ = time_operation(lambda: list(db[USERS_COLLECTION].aggregate(pipeline)))
    return ms


def bench_delete_one_user(db):
    """Delete a single user by user_id."""
    target = random.randint(0, NUM_USERS - 1)
    ms, _ = time_operation(
        db[USERS_COLLECTION].delete_one, {"user_id": target}
    )
    return ms


def bench_delete_by_status(db):
    """Delete all cancelled orders."""
    ms, _ = time_operation(
        db[ORDERS_COLLECTION].delete_many, {"status": "cancelled"}
    )
    return ms


def bench_delete_by_amount(db):
    """Delete orders with total_amount < 20."""
    ms, _ = time_operation(
        db[ORDERS_COLLECTION].delete_many, {"total_amount": {"$lt": 20}}
    )
    return ms


def bench_delete_by_user_and_status(db):
    """Delete pending orders for a specific user."""
    target = random.randint(0, NUM_USERS - 1)
    ms, _ = time_operation(
        db[ORDERS_COLLECTION].delete_many, {"user_id": target, "status": "pending"}
    )
    return ms


def bench_delete_all(db):
    """Delete all users and orders. Returns ms."""
    ms_users, _ = time_operation(db[USERS_COLLECTION].delete_many, {})
    ms_orders, _ = time_operation(db[ORDERS_COLLECTION].delete_many, {})
    return ms_users + ms_orders


# ---------------------------------------------------------------------------
# Run all benchmarks for one target
# ---------------------------------------------------------------------------
BENCHMARKS_PHASE1 = [
    ("Insert (bulk)",              bench_insert),
    # -- Updates --
    ("Update: inc age",            bench_update_inc),
    ("Update: set nested",         bench_update_set_nested),
    ("Update: push array",         bench_update_push_array),
    ("Update: conditional",        bench_update_conditional),
    ("Update: orders status",      bench_update_orders_status),
    # -- Finds --
    ("Find: by user_id",           bench_find_by_id),
    ("Find: age range",            bench_find_range),
    ("Find: regex email",          bench_find_regex),
    ("Find: projection",           bench_find_projection),
    ("Find: sort + limit",         bench_find_sort_limit),
    ("Find: orders for user",      bench_find_orders_for_user),
    ("Find: multi-field",          bench_find_multi_field),
    ("Find: in array",             bench_find_in_array),
    ("Find: orders by amount",     bench_find_orders_by_amount),
    ("Find: orders status+user",   bench_find_orders_status_user),
    ("Find: count by status",      bench_find_count),
    ("Find: distinct status",      bench_find_distinct),
    # -- Aggregations --
    ("Agg: group by status",       bench_agg_group_status),
    ("Agg: unwind items",          bench_agg_unwind_items),
    ("Agg: bucket amounts",        bench_agg_bucket_amount),
    ("Agg: user order stats",      bench_agg_user_order_stats),
    ("Agg: date breakdown",        bench_agg_date_breakdown),
    # -- Lookup --
    ("Lookup (users→orders)",      bench_lookup),
    # -- Deletes --
    ("Delete: one user",           bench_delete_one_user),
    ("Delete: by status",          bench_delete_by_status),
    ("Delete: by amount",          bench_delete_by_amount),
    ("Delete: user+status",        bench_delete_by_user_and_status),
    ("Delete: all",                bench_delete_all),
]

# Benchmarks to re-run after adding optimized indexes
BENCHMARKS_PHASE2 = [
    ("Update: inc age",            bench_update_inc),
    ("Update: set nested",         bench_update_set_nested),
    ("Update: push array",         bench_update_push_array),
    ("Update: conditional",        bench_update_conditional),
    ("Update: orders status",      bench_update_orders_status),
    ("Find: by user_id",           bench_find_by_id),
    ("Find: age range",            bench_find_range),
    ("Find: regex email",          bench_find_regex),
    ("Find: projection",           bench_find_projection),
    ("Find: sort + limit",         bench_find_sort_limit),
    ("Find: orders for user",      bench_find_orders_for_user),
    ("Find: multi-field",          bench_find_multi_field),
    ("Find: in array",             bench_find_in_array),
    ("Find: orders by amount",     bench_find_orders_by_amount),
    ("Find: orders status+user",   bench_find_orders_status_user),
    ("Find: count by status",      bench_find_count),
    ("Find: distinct status",      bench_find_distinct),
    ("Agg: group by status",       bench_agg_group_status),
    ("Agg: unwind items",          bench_agg_unwind_items),
    ("Agg: bucket amounts",        bench_agg_bucket_amount),
    ("Agg: user order stats",      bench_agg_user_order_stats),
    ("Agg: date breakdown",        bench_agg_date_breakdown),
    ("Lookup (users→orders)",      bench_lookup),
    ("Delete: one user",           bench_delete_one_user),
    ("Delete: by status",          bench_delete_by_status),
    ("Delete: by amount",          bench_delete_by_amount),
    ("Delete: user+status",        bench_delete_by_user_and_status),
    ("Delete: all",                bench_delete_all),
]


def create_optimized_indexes(db):
    """Create indexes optimized for each benchmark operation."""
    users = db[USERS_COLLECTION]
    orders = db[ORDERS_COLLECTION]

    # -- Users single-field indexes --
    users.create_index("age")                  # Find: age range, sort+limit
    users.create_index("email")                # Find: regex email
    users.create_index("status")               # Delete: one user uses user_id (already indexed)
    users.create_index("tags")                 # Find: in array (multikey index)
    users.create_index("preferences.language")  # Find: multi-field filter

    # -- Users compound indexes --
    users.create_index([("age", 1), ("preferences.language", 1)])  # Find: multi-field (age+lang)
    users.create_index([("age", -1), ("user_id", 1)])              # Find: sort+limit by age desc
    users.create_index([("status", 1), ("age", 1)])                # Update: conditional (age>40)
    users.create_index([("email", 1), ("user_id", 1)])             # Find: regex email + projection

    # -- Orders single-field indexes --
    orders.create_index("status")              # Find: count, Delete: by status, Update: status
    orders.create_index("total_amount")        # Find: orders by amount, Delete: by amount
    orders.create_index("created_at")          # Agg: date breakdown

    # -- Orders compound indexes --
    orders.create_index([("user_id", 1), ("status", 1)])            # Find: orders status+user, Delete: user+status
    orders.create_index([("user_id", 1), ("total_amount", 1)])      # Agg: user order stats
    orders.create_index([("status", 1), ("total_amount", 1)])       # Agg: group by status, Update: orders status
    orders.create_index([("total_amount", -1)])                     # Find: orders by amount desc + limit
    orders.create_index([("created_at", 1), ("total_amount", 1)])   # Agg: date breakdown


def _run_phase(db, benchmarks, label, phase_name):
    """Run a list of benchmarks and return {name: {min, max, avg, median}}."""
    print(f"\n  --- {phase_name} ---")
    results: dict[str, dict] = {}

    for name, func in benchmarks:
        times = []
        # Insert runs only once; all other operations use ITERATIONS
        iters = 1 if name == "Insert (bulk)" else ITERATIONS
        for iteration in range(iters):
            # For insert: clean first so we can re-insert fresh data
            if name == "Insert (bulk)":
                cleanup(db)
                db[USERS_COLLECTION].create_index("user_id", unique=True)
                db[ORDERS_COLLECTION].create_index("user_id")
                db[ORDERS_COLLECTION].create_index("order_id", unique=True)

            # Ensure data exists (may have been wiped by a prior delete benchmark)
            if name != "Insert (bulk)" and db[USERS_COLLECTION].count_documents({}) == 0:
                bench_insert(db)
            if name != "Insert (bulk)" and db[ORDERS_COLLECTION].count_documents({}) == 0:
                bench_insert(db)

            ms = func(db)
            times.append(ms)
            print(f"    {name} [iter {iteration+1}] => {ms:>9.2f} ms")

        results[name] = {
            "min": min(times),
            "max": max(times),
            "avg": statistics.mean(times),
            "median": statistics.median(times),
        }

    return results


def run_benchmarks(uri: str, label: str) -> tuple[dict, dict]:
    """Connect, run Phase 1 (basic indexes), add optimized indexes, run Phase 2."""
    print(f"\n{'='*60}")
    print(f"  Connecting to {label} ...")
    print(f"{'='*60}")

    client = MongoClient(uri, serverSelectionTimeoutMS=10000)
    try:
        client.admin.command("ping")
        print(f"  ✓ Connected to {label}")
    except ConnectionFailure as exc:
        print(f"  ✗ Cannot connect to {label}: {exc}")
        return {}, {}

    db = client[DB_NAME]

    # Clean up any leftover data from previous runs
    cleanup(db)

    # Create basic indexes only
    db[USERS_COLLECTION].create_index("user_id", unique=True)
    db[ORDERS_COLLECTION].create_index("user_id")
    db[ORDERS_COLLECTION].create_index("order_id", unique=True)

    # ---- Phase 1: basic indexes ----
    phase1 = _run_phase(db, BENCHMARKS_PHASE1, label, "Phase 1: Basic indexes")

    # ---- Add optimized indexes ----
    print(f"\n  >> Creating optimized single-field & compound indexes ...")
    create_optimized_indexes(db)
    print(f"  >> Indexes created. Re-running find / update / aggregation benchmarks ...\n")

    # Ensure data exists for Phase 2
    if db[USERS_COLLECTION].count_documents({}) == 0:
        bench_insert(db)

    # ---- Phase 2: optimized indexes ----
    phase2 = _run_phase(db, BENCHMARKS_PHASE2, label, "Phase 2: Optimized indexes")

    # Final cleanup
    cleanup(db)
    client.close()
    return phase1, phase2


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def _print_comparison(cosmos_results, atlas_results, benchmarks, title):
    """Print a side-by-side comparison table."""
    print("\n\n" + "=" * 90)
    print(f"  {title}  (all times in milliseconds)")
    print("=" * 90)

    headers = [
        "Operation",
        "Cosmos min", "Cosmos avg", "Cosmos max",
        "Atlas min", "Atlas avg", "Atlas max",
        "Diff %", "Winner",
    ]
    rows = []
    for name, _ in benchmarks:
        c = cosmos_results.get(name, {})
        a = atlas_results.get(name, {})
        c_avg = c.get("avg", float("inf"))
        a_avg = a.get("avg", float("inf"))
        winner = "Cosmos DB" if c_avg <= a_avg else "Atlas"

        if c_avg and a_avg and c_avg != float("inf") and a_avg != float("inf"):
            faster = min(c_avg, a_avg)
            slower = max(c_avg, a_avg)
            diff_pct = ((slower - faster) / faster) * 100
            diff_str = f"{diff_pct:+.1f}%"
        else:
            diff_str = "N/A"

        rows.append([
            name,
            f"{c.get('min', 0):.2f}" if c else "N/A",
            f"{c_avg:.2f}" if c else "N/A",
            f"{c.get('max', 0):.2f}" if c else "N/A",
            f"{a.get('min', 0):.2f}" if a else "N/A",
            f"{a_avg:.2f}" if a else "N/A",
            f"{a.get('max', 0):.2f}" if a else "N/A",
            diff_str,
            winner,
        ])

    print(tabulate(rows, headers=headers, tablefmt="fancy_grid"))


def _print_index_impact(phase1, phase2, label, benchmarks):
    """Show before/after per operation with gain or loss for one target."""
    print(f"\n{'='*90}")
    print(f"  INDEX IMPACT — {label}  (before vs after optimized indexes)")
    print(f"{'='*90}")
    headers = [
        "Operation",
        "No Idx min", "No Idx avg", "No Idx max",
        "Idx min", "Idx avg", "Idx max",
        "Gain/Loss", "Verdict",
    ]
    rows = []
    for name, _ in benchmarks:
        b = phase1.get(name, {})
        a = phase2.get(name, {})
        b_avg = b.get("avg")
        a_avg = a.get("avg")
        if b_avg and a_avg and b_avg > 0:
            pct = ((b_avg - a_avg) / b_avg) * 100
            gain_str = f"{pct:+.1f}%"
            verdict = "FASTER" if pct > 1 else ("SLOWER" if pct < -1 else "~SAME")
        else:
            gain_str = "N/A"
            verdict = "N/A"
        rows.append([
            name,
            f"{b.get('min', 0):.2f}" if b else "N/A",
            f"{b_avg:.2f}" if b_avg else "N/A",
            f"{b.get('max', 0):.2f}" if b else "N/A",
            f"{a.get('min', 0):.2f}" if a else "N/A",
            f"{a_avg:.2f}" if a_avg else "N/A",
            f"{a.get('max', 0):.2f}" if a else "N/A",
            gain_str,
            verdict,
        ])
    print(tabulate(rows, headers=headers, tablefmt="fancy_grid"))


def _print_combined_index_impact(cosmos_p1, cosmos_p2, atlas_p1, atlas_p2, benchmarks):
    """Side-by-side index impact comparison: Cosmos DB vs Atlas."""
    print(f"\n{'='*90}")
    print(f"  COMBINED INDEX IMPACT — Cosmos DB vs Atlas")
    print(f"{'='*90}")
    headers = [
        "Operation",
        "Cosmos No Idx", "Cosmos w/ Idx", "Cosmos Gain",
        "Atlas No Idx", "Atlas w/ Idx", "Atlas Gain",
        "Best Indexed",
    ]
    rows = []
    for name, _ in benchmarks:
        cb = cosmos_p1.get(name, {}).get("avg")
        ca = cosmos_p2.get(name, {}).get("avg")
        ab = atlas_p1.get(name, {}).get("avg")
        aa = atlas_p2.get(name, {}).get("avg")

        c_gain = ((cb - ca) / cb * 100) if cb and ca and cb > 0 else None
        a_gain = ((ab - aa) / ab * 100) if ab and aa and ab > 0 else None

        if ca and aa:
            best = "Cosmos DB" if ca <= aa else "Atlas"
        else:
            best = "N/A"

        rows.append([
            name,
            f"{cb:.2f}" if cb else "N/A",
            f"{ca:.2f}" if ca else "N/A",
            f"{c_gain:+.1f}%" if c_gain is not None else "N/A",
            f"{ab:.2f}" if ab else "N/A",
            f"{aa:.2f}" if aa else "N/A",
            f"{a_gain:+.1f}%" if a_gain is not None else "N/A",
            best,
        ])
    print(tabulate(rows, headers=headers, tablefmt="fancy_grid"))


# ---------------------------------------------------------------------------
# Word document export
# ---------------------------------------------------------------------------
def _add_table_to_doc(doc, headers, rows, title):
    """Add a styled table with a heading to the Word document."""
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph("")


def _build_comparison_rows(cosmos_results, atlas_results, benchmarks):
    headers = [
        "Operation",
        "Cosmos min", "Cosmos avg", "Cosmos max",
        "Atlas min", "Atlas avg", "Atlas max",
        "Diff %", "Winner",
    ]
    rows = []
    for name, _ in benchmarks:
        c = cosmos_results.get(name, {})
        a = atlas_results.get(name, {})
        c_avg = c.get("avg", float("inf"))
        a_avg = a.get("avg", float("inf"))
        winner = "Cosmos DB" if c_avg <= a_avg else "Atlas"
        if c_avg and a_avg and c_avg != float("inf") and a_avg != float("inf"):
            faster = min(c_avg, a_avg)
            slower = max(c_avg, a_avg)
            diff_str = f"{((slower - faster) / faster) * 100:+.1f}%"
        else:
            diff_str = "N/A"
        rows.append([
            name,
            f"{c.get('min', 0):.2f}" if c else "N/A",
            f"{c_avg:.2f}" if c else "N/A",
            f"{c.get('max', 0):.2f}" if c else "N/A",
            f"{a.get('min', 0):.2f}" if a else "N/A",
            f"{a_avg:.2f}" if a else "N/A",
            f"{a.get('max', 0):.2f}" if a else "N/A",
            diff_str, winner,
        ])
    return headers, rows


def _build_index_impact_rows(phase1, phase2, benchmarks):
    headers = [
        "Operation",
        "No Idx min", "No Idx avg", "No Idx max",
        "Idx min", "Idx avg", "Idx max",
        "Gain/Loss", "Verdict",
    ]
    rows = []
    for name, _ in benchmarks:
        b = phase1.get(name, {})
        a = phase2.get(name, {})
        b_avg = b.get("avg")
        a_avg = a.get("avg")
        if b_avg and a_avg and b_avg > 0:
            pct = ((b_avg - a_avg) / b_avg) * 100
            gain_str = f"{pct:+.1f}%"
            verdict = "FASTER" if pct > 1 else ("SLOWER" if pct < -1 else "~SAME")
        else:
            gain_str = "N/A"
            verdict = "N/A"
        rows.append([
            name,
            f"{b.get('min', 0):.2f}" if b else "N/A",
            f"{b_avg:.2f}" if b_avg else "N/A",
            f"{b.get('max', 0):.2f}" if b else "N/A",
            f"{a.get('min', 0):.2f}" if a else "N/A",
            f"{a_avg:.2f}" if a_avg else "N/A",
            f"{a.get('max', 0):.2f}" if a else "N/A",
            gain_str, verdict,
        ])
    return headers, rows


def _build_combined_impact_rows(cosmos_p1, cosmos_p2, atlas_p1, atlas_p2, benchmarks):
    headers = [
        "Operation",
        "Cosmos No Idx", "Cosmos w/ Idx", "Cosmos Gain",
        "Atlas No Idx", "Atlas w/ Idx", "Atlas Gain",
        "Best Indexed",
    ]
    rows = []
    for name, _ in benchmarks:
        cb = cosmos_p1.get(name, {}).get("avg")
        ca = cosmos_p2.get(name, {}).get("avg")
        ab = atlas_p1.get(name, {}).get("avg")
        aa = atlas_p2.get(name, {}).get("avg")
        c_gain = ((cb - ca) / cb * 100) if cb and ca and cb > 0 else None
        a_gain = ((ab - aa) / ab * 100) if ab and aa and ab > 0 else None
        best = ("Cosmos DB" if ca <= aa else "Atlas") if ca and aa else "N/A"
        rows.append([
            name,
            f"{cb:.2f}" if cb else "N/A",
            f"{ca:.2f}" if ca else "N/A",
            f"{c_gain:+.1f}%" if c_gain is not None else "N/A",
            f"{ab:.2f}" if ab else "N/A",
            f"{aa:.2f}" if aa else "N/A",
            f"{a_gain:+.1f}%" if a_gain is not None else "N/A",
            best,
        ])
    return headers, rows


def export_to_word(cosmos_p1, cosmos_p2, atlas_p1, atlas_p2, filename="benchmark_results.docx"):
    """Export all benchmark results to a Word document."""
    doc = Document()

    # Title
    title = doc.add_heading("MongoDB Benchmark Results", level=0)
    doc.add_paragraph(
        f"Cosmos DB (DocumentDB) vs MongoDB Atlas\n"
        f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"Users: {NUM_USERS}  |  Orders/user: {ORDERS_PER_USER}  |  Iterations: {ITERATIONS}"
    )

    # Phase 1 comparison
    h, r = _build_comparison_rows(cosmos_p1, atlas_p1, BENCHMARKS_PHASE1)
    _add_table_to_doc(doc, h, r, "Phase 1 — Basic Indexes")

    # Phase 2 comparison
    h, r = _build_comparison_rows(cosmos_p2, atlas_p2, BENCHMARKS_PHASE2)
    _add_table_to_doc(doc, h, r, "Phase 2 — Optimized Indexes (single-field + compound)")

    # Index impact per target
    h, r = _build_index_impact_rows(cosmos_p1, cosmos_p2, BENCHMARKS_PHASE2)
    _add_table_to_doc(doc, h, r, "Index Impact — Cosmos DB")

    h, r = _build_index_impact_rows(atlas_p1, atlas_p2, BENCHMARKS_PHASE2)
    _add_table_to_doc(doc, h, r, "Index Impact — MongoDB Atlas")

    # Combined index impact
    h, r = _build_combined_impact_rows(cosmos_p1, cosmos_p2, atlas_p1, atlas_p2, BENCHMARKS_PHASE2)
    _add_table_to_doc(doc, h, r, "Combined Index Impact — Cosmos DB vs Atlas")

    # Detail tables
    for label, p1, p2 in [("Cosmos DB", cosmos_p1, cosmos_p2), ("Atlas", atlas_p1, atlas_p2)]:
        for phase_label, results in [("Phase 1 (basic)", p1), ("Phase 2 (optimized)", p2)]:
            detail_h = ["Operation", "Min", "Max", "Avg", "Median"]
            detail_r = []
            for name, stats in results.items():
                detail_r.append([
                    name,
                    f"{stats['min']:.2f}",
                    f"{stats['max']:.2f}",
                    f"{stats['avg']:.2f}",
                    f"{stats['median']:.2f}",
                ])
            _add_table_to_doc(doc, detail_h, detail_r, f"{label} — {phase_label} Detail")

    doc.save(filename)
    print(f"\n  >> Results saved to {filename}")


def main():
    print("\n" + "=" * 60)
    print("  MongoDB Benchmark: Cosmos DB (DocumentDB) vs Atlas")
    print(f"  Users: {NUM_USERS}  |  Orders/user: {ORDERS_PER_USER}  |  Iterations: {ITERATIONS}")
    print("=" * 60)

    cosmos_p1, cosmos_p2 = run_benchmarks(COSMOSDB_URI, "Azure Cosmos DB (DocumentDB)")
    atlas_p1, atlas_p2 = run_benchmarks(ATLAS_URI, "MongoDB Atlas")

    # ---- Phase 1 comparison (basic indexes) ----
    _print_comparison(cosmos_p1, atlas_p1, BENCHMARKS_PHASE1,
                      "PHASE 1 — BASIC INDEXES")

    # ---- Phase 2 comparison (optimized indexes) ----
    _print_comparison(cosmos_p2, atlas_p2, BENCHMARKS_PHASE2,
                      "PHASE 2 — OPTIMIZED INDEXES (single-field + compound)")

    # ---- Index impact per target ----
    _print_index_impact(cosmos_p1, cosmos_p2, "Cosmos DB", BENCHMARKS_PHASE2)
    _print_index_impact(atlas_p1, atlas_p2, "MongoDB Atlas", BENCHMARKS_PHASE2)

    # ---- Combined side-by-side index impact ----
    _print_combined_index_impact(cosmos_p1, cosmos_p2, atlas_p1, atlas_p2, BENCHMARKS_PHASE2)

    # ---- Detailed per-target tables ----
    for label, p1, p2 in [
        ("Cosmos DB", cosmos_p1, cosmos_p2),
        ("Atlas", atlas_p1, atlas_p2),
    ]:
        for phase_label, results in [("Phase 1 (basic)", p1), ("Phase 2 (optimized)", p2)]:
            print(f"\n--- {label} — {phase_label} detail ---")
            detail_rows = []
            for name, stats in results.items():
                detail_rows.append([
                    name,
                    f"{stats['min']:.2f}",
                    f"{stats['max']:.2f}",
                    f"{stats['avg']:.2f}",
                    f"{stats['median']:.2f}",
                ])
            print(tabulate(detail_rows, headers=["Operation", "Min", "Max", "Avg", "Median"],
                           tablefmt="fancy_grid"))

    # ---- Export to Word ----
    export_to_word(cosmos_p1, cosmos_p2, atlas_p1, atlas_p2)


if __name__ == "__main__":
    main()
